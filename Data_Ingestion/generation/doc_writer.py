"""
Document Writer
================
Assembles generated section versions into a final output document.

Supported output formats:
  - Markdown (.md)      — trivial assembly, full fidelity
  - Word (.docx)        — python-docx with Markdown parsing (headings, bullets, tables, bold)
  - PDF                 — requires weasyprint (optional; falls back to Markdown if not installed)

Output files are saved to:
  local_storage/outputs/{job_id}/{filename}      (LOCAL_DB=true)
  GCS: gs://{bucket}/outputs/{job_id}/{filename} (LOCAL_DB=false)

The export_job() function is the main entry point — it pulls the latest accepted
(or current) version of each section and assembles the document.
"""

from __future__ import annotations
import logging
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

LOCAL_DB = os.environ.get("LOCAL_DB", "true").lower() == "true"


# ─────────────────────────────────────────────────────────────────────────────
# Main export entry point
# ─────────────────────────────────────────────────────────────────────────────

def export_job(job_id: str, output_format: Optional[str] = None) -> tuple[Path, str]:
    """
    Assemble the final document for a job from accepted/current section versions.

    Args:
        job_id:        The generation job ID
        output_format: Override format — 'Word (.docx)', 'PDF', or 'Markdown'.
                       If None, uses the format stored on the job.

    Returns:
        (file_path, mime_type) — local path to the exported file and its MIME type
    """
    from generation.db import GenerationJob, Section, get_session

    with get_session() as session:
        job = session.get(GenerationJob, job_id)
        if not job:
            raise ValueError(f"Job {job_id} not found")

        fmt           = output_format or job.output_format or "Markdown"
        doc_type      = job.document_type
        project_name  = _extract_project_name(job.user_inputs_json)

        # Collect sections in order — prefer accepted version, fall back to current
        sections_content: list[dict] = []
        for sec in sorted(job.sections, key=lambda s: s.order_index):
            if not sec.versions:
                continue
            # Pick accepted version first, then latest
            accepted = next((v for v in sec.versions if v.is_accepted), None)
            latest   = max(sec.versions, key=lambda v: v.version_number)
            chosen   = accepted or latest

            sections_content.append({
                "key":     getattr(sec, "section_key", None) or "",
                "title":   sec.section_title,
                "content": chosen.content,
                "version": chosen.version_number,
            })

    if not sections_content:
        raise ValueError("No completed sections to export")

    # Build output directory
    out_dir = _output_dir(job_id)
    out_dir.mkdir(parents=True, exist_ok=True)

    safe_name = re.sub(r"[^\w\s-]", "", project_name)[:40].strip().replace(" ", "_")
    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M")

    if "word" in fmt.lower() or "docx" in fmt.lower():
        filename  = f"{safe_name}_{timestamp}.docx"
        out_path  = out_dir / filename
        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        _write_styled_docx(out_path, doc_type, project_name, sections_content)

    elif "pdf" in fmt.lower():
        pdf_path  = out_dir / f"{safe_name}_{timestamp}.pdf"
        try:
            _write_pdf(pdf_path, doc_type, project_name, sections_content)
            out_path  = pdf_path
            mime_type = "application/pdf"
        except _NoPdfConverter as exc:
            # No PDF library installed — serve DOCX instead so the user still gets a file
            logger.warning(
                "PDF export unavailable (%s) — falling back to DOCX. "
                "Install docx2pdf, weasyprint, or xhtml2pdf to enable PDF export.", exc
            )
            filename  = f"{safe_name}_{timestamp}.docx"
            out_path  = out_dir / filename
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            _write_styled_docx(out_path, doc_type, project_name, sections_content)
        filename = out_path.name

    else:   # Markdown (default)
        filename  = f"{safe_name}_{timestamp}.md"
        out_path  = out_dir / filename
        mime_type = "text/markdown"
        _write_markdown(out_path, doc_type, project_name, sections_content)

    logger.info("Exported %s (%d sections) → %s", job_id, len(sections_content), out_path)
    return out_path, mime_type


# ─────────────────────────────────────────────────────────────────────────────
# Format writers
# ─────────────────────────────────────────────────────────────────────────────

def _write_markdown(
    out_path: Path,
    doc_type: str,
    project_name: str,
    sections: list[dict],
) -> None:
    lines = [
        f"# {doc_type}",
        f"**Project:** {project_name}",
        f"**Generated:** {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}",
        "",
        "---",
        "",
    ]
    for sec in sections:
        lines.append(f"## {sec['title']}")
        lines.append("")
        lines.append(sec["content"])
        lines.append("")
        lines.append("---")
        lines.append("")

    out_path.write_text("\n".join(lines), encoding="utf-8")


def _write_styled_docx(
    out_path: Path,
    doc_type: str,
    project_name: str,
    sections: list[dict],
) -> None:
    """
    Route to the Adani-styled Word formatter. Document structure now comes from
    the expert-mapping-compiled templates (single source of truth), so BRD renders
    through the GENERIC structured formatter like every other type — its numbered
    sections and sub-headings come straight from the generated section content.
    NIT keeps its Section-I/II/III/X layout (static boilerplate keyed by section).
    All paths share the same cover, header+logo, footer, styled headings, and
    dark-blue tables — so every document type looks consistent.
    """
    if _is_nit(doc_type):
        from generation.brd_formatter import format_nit_docx
        sections_by_key = {s["key"]: s["content"] for s in sections if s.get("key")}
        format_nit_docx(sections_by_key, project_name, out_path)
    else:
        from generation.brd_formatter import format_structured_docx
        format_structured_docx(doc_type, project_name, sections, out_path)


def _write_docx(
    out_path: Path,
    doc_type: str,
    project_name: str,
    sections: list[dict],
) -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Document title
    title_para = doc.add_heading(doc_type, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_para.add_run(
        f"Project: {project_name}  |  Generated: {datetime.utcnow().strftime('%Y-%m-%d')}"
    )
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()   # spacer

    for sec in sections:
        # Section heading
        doc.add_heading(sec["title"], level=1)
        # Parse and render Markdown content
        _render_markdown_to_docx(doc, sec["content"])
        doc.add_paragraph()   # spacer between sections

    doc.save(str(out_path))


def _write_pdf(
    out_path: Path,
    doc_type: str,
    project_name: str,
    sections: list[dict],
) -> Path:
    """
    Convert content to PDF.  Conversion chain (first available wins):
      1. docx2pdf  — styled .docx via Microsoft Word COM (Windows + Office). Highest
                     fidelity (full Adani cover/logo/footer), used on a local Windows
                     box with Office; unavailable on Databricks.
      2. xhtml2pdf — pure-Python HTML → PDF, styled from generation/doc_style.py so it
                     matches the preview + .docx brand (dark-blue tables, headings).
                     This is the cross-platform / Databricks path. No LibreOffice.

    If neither is available this raises _NoPdfConverter so the caller falls back
    to serving the DOCX instead. (LibreOffice was removed from the backend
    2026-07-17 — see doc_style.py.)

    Returns the actual output path (always `out_path` on success).
    """
    # ── Method 1: docx2pdf (Word COM, Windows with MS Office) ────────────────
    try:
        from docx2pdf import convert as _docx2pdf
        docx_tmp = out_path.with_suffix(".docx")
        _write_styled_docx(docx_tmp, doc_type, project_name, sections)
        _docx2pdf(str(docx_tmp), str(out_path))
        try:
            docx_tmp.unlink()
        except OSError:
            pass
        logger.info("PDF written via docx2pdf → %s", out_path)
        return out_path
    except ImportError:
        pass   # docx2pdf not installed (e.g. Databricks)
    except Exception as e:
        logger.warning("docx2pdf failed: %s — trying xhtml2pdf", e)

    # ── Build styled HTML (brand palette from doc_style — matches preview/.docx) ──
    def _build_html() -> str:
        from generation.doc_style import pdf_css
        body_lines = [
            '<div class="id-rule">',
            f'<p class="id-title">{doc_type}</p>',
            f'<p class="id-meta"><strong>Project:</strong> {project_name} &nbsp;|&nbsp; '
            f'<strong>Generated:</strong> {datetime.utcnow().strftime("%Y-%m-%d")}</p>',
            '</div>',
        ]
        for i, sec in enumerate(sections, start=1):
            body_lines.append(f"<h2>{i}. {sec['title']}</h2>")
            body_lines.append(_md_to_html(sec["content"]))
        return (
            f"<!DOCTYPE html><html><head><meta charset='UTF-8'>"
            f"<style>{pdf_css()}</style></head><body>{''.join(body_lines)}</body></html>"
        )

    # ── Method 2: xhtml2pdf (pure Python — the Databricks path) ──────────────
    try:
        from xhtml2pdf import pisa
        with open(out_path, "wb") as fh:
            result = pisa.CreatePDF(src=_build_html(), dest=fh)
        if result.err:
            raise RuntimeError(f"xhtml2pdf reported errors: {result.err}")
        logger.info("PDF written via xhtml2pdf → %s", out_path)
        return out_path
    except ImportError:
        pass

    # ── No converter available ───────────────────────────────────────────────
    raise _NoPdfConverter(
        "No PDF converter is installed. Install xhtml2pdf (cross-platform) "
        "or docx2pdf (needs MS Word on Windows)."
    )


class _NoPdfConverter(RuntimeError):
    """Raised when no PDF conversion library is available."""


# ─────────────────────────────────────────────────────────────────────────────
# Markdown → python-docx renderer
# ─────────────────────────────────────────────────────────────────────────────

def _render_markdown_to_docx(doc, md_text: str) -> None:
    """
    Parse a subset of Markdown and add formatted content to a python-docx Document.
    Handles: headings (##, ###), bold (**text**), bullets (- item), numbered lists,
    horizontal rules, code blocks, and plain paragraphs.
    Tables are rendered as python-docx tables.
    """
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    lines     = md_text.split("\n")
    i         = 0
    in_code   = False
    code_buf: list[str] = []
    table_buf: list[str] = []

    while i < len(lines):
        line = lines[i]

        # ── Code block ─────────────────────────────────────────────────────
        if line.strip().startswith("```"):
            if in_code:
                # End of code block
                if code_buf:
                    p = doc.add_paragraph()
                    style = _safe_style(doc, "No Spacing", "Normal")
                    if style:
                        p.style = style
                    run = p.add_run("\n".join(code_buf))
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                code_buf = []
                in_code  = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # ── Table detection (| col | col |) ───────────────────────────────
        if line.strip().startswith("|"):
            table_buf.append(line)
            i += 1
            # Collect all consecutive table lines
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_buf.append(lines[i])
                i += 1
            _add_table_to_docx(doc, table_buf)
            table_buf = []
            continue

        # ── Headings ───────────────────────────────────────────────────────
        if line.startswith("### "):
            doc.add_heading(_strip_inline_md(line[4:]), level=3)
        elif line.startswith("## "):
            doc.add_heading(_strip_inline_md(line[3:]), level=2)
        elif line.startswith("# "):
            doc.add_heading(_strip_inline_md(line[2:]), level=1)

        # ── Horizontal rule ────────────────────────────────────────────────
        elif line.strip() in ("---", "***", "___"):
            doc.add_paragraph("─" * 60)

        # ── Bullet list ────────────────────────────────────────────────────
        elif line.startswith("- ") or line.startswith("* "):
            _add_run_with_inline_md(doc.add_paragraph(style="List Bullet"), line[2:])

        # ── Numbered list ──────────────────────────────────────────────────
        elif re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s", "", line)
            _add_run_with_inline_md(doc.add_paragraph(style="List Number"), text)

        # ── Sub-bullet (  - or    -) ───────────────────────────────────────
        elif re.match(r"^\s{2,}- ", line):
            _add_run_with_inline_md(doc.add_paragraph(style="List Bullet 2"), line.strip()[2:])

        # ── Empty line ─────────────────────────────────────────────────────
        elif not line.strip():
            pass  # skip — paragraph spacing handles visual gaps

        # ── Normal paragraph ───────────────────────────────────────────────
        else:
            _add_run_with_inline_md(doc.add_paragraph(), line)

        i += 1


def _safe_style(doc, *names):
    """
    Return the first available python-docx Style by name, or None.
    python-docx's Styles object is NOT a dict — it has no .get() — so we
    probe each name with a try/except KeyError.
    """
    for name in names:
        try:
            return doc.styles[name]
        except KeyError:
            continue
    return None


def _add_table_to_docx(doc, table_lines: list[str]) -> None:
    """Convert Markdown table lines to a python-docx Table."""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    rows = []
    for raw in table_lines:
        if re.match(r"^\|[-| :]+\|$", raw.strip()):
            continue  # separator row
        cells = [c.strip() for c in raw.strip().strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return

    max_cols = max(len(r) for r in rows)
    tbl      = doc.add_table(rows=len(rows), cols=max_cols)
    style    = _safe_style(doc, "Table Grid", "Normal Table", "Table Normal")
    if style:
        tbl.style = style

    for r_idx, row_cells in enumerate(rows):
        for c_idx in range(max_cols):
            cell_text = row_cells[c_idx] if c_idx < len(row_cells) else ""
            cell = tbl.rows[r_idx].cells[c_idx]
            para = cell.paragraphs[0]
            run  = para.add_run(_strip_inline_md(cell_text))
            if r_idx == 0:
                run.bold = True


def _add_run_with_inline_md(para, text: str) -> None:
    """
    Add runs to a paragraph, handling **bold**, *italic*, and `code` spans.
    """
    # Pattern to split on inline formatting tokens
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    parts   = pattern.split(text)

    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run      = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run        = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            from docx.shared import Pt
            run           = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            para.add_run(part)


def _strip_inline_md(text: str) -> str:
    """Remove inline MD markers for contexts that don't support formatting."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*",     r"\1", text)
    text = re.sub(r"`(.+?)`",       r"\1", text)
    return text.strip()


def _md_to_html(md: str) -> str:
    """
    Minimal Markdown → HTML converter used by the PDF export path.
    Handles: headings, bold/italic/code spans, bullet lists, numbered lists,
    horizontal rules, fenced code blocks, and pipe tables.
    Falls back to <pre> wrapping if markdown2 / mistune is available.
    """
    # Try markdown2 first (optional dep, best fidelity)
    try:
        import markdown2
        return markdown2.markdown(md, extras=["tables", "fenced-code-blocks"])
    except ImportError:
        pass

    # Try mistune (another optional dep)
    try:
        import mistune
        return mistune.html(md)
    except ImportError:
        pass

    # Hand-rolled minimal converter
    import html as _html
    lines = md.split("\n")
    out: list[str] = []
    in_code  = False
    in_ul    = False
    in_ol    = False

    def close_lists():
        nonlocal in_ul, in_ol
        if in_ul:
            out.append("</ul>"); in_ul = False
        if in_ol:
            out.append("</ol>"); in_ol = False

    for line in lines:
        # Code fence
        if line.strip().startswith("```"):
            close_lists()
            if in_code:
                out.append("</code></pre>"); in_code = False
            else:
                out.append("<pre><code>"); in_code = True
            continue
        if in_code:
            out.append(_html.escape(line))
            continue

        # Headings
        if line.startswith("### "):
            close_lists(); out.append(f"<h3>{_inline(line[4:])}</h3>"); continue
        if line.startswith("## "):
            close_lists(); out.append(f"<h2>{_inline(line[3:])}</h2>"); continue
        if line.startswith("# "):
            close_lists(); out.append(f"<h1>{_inline(line[2:])}</h1>"); continue

        # HR
        if line.strip() in ("---", "***", "___"):
            close_lists(); out.append("<hr>"); continue

        # Bullet
        if line.startswith("- ") or line.startswith("* "):
            if not in_ul:
                close_lists(); out.append("<ul>"); in_ul = True
            out.append(f"<li>{_inline(line[2:])}</li>"); continue

        # Numbered
        m = re.match(r"^(\d+)\.\s(.+)", line)
        if m:
            if not in_ol:
                close_lists(); out.append("<ol>"); in_ol = True
            out.append(f"<li>{_inline(m.group(2))}</li>"); continue

        # Pipe table row
        if line.strip().startswith("|"):
            close_lists()
            if re.match(r"^\|[-| :]+\|$", line.strip()):
                continue  # separator row
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            tag = "th" if not out or not out[-1].startswith("<tr") else "td"
            out.append("<tr>" + "".join(f"<{tag}>{_inline(c)}</{tag}>" for c in cells) + "</tr>")
            continue

        # Empty line
        if not line.strip():
            close_lists(); out.append(""); continue

        # Normal paragraph
        close_lists()
        out.append(f"<p>{_inline(line)}</p>")

    close_lists()
    # Wrap table rows in <table>
    result = "\n".join(out)
    result = re.sub(r"(<tr>.*?</tr>\n?)+", lambda m: f"<table>{m.group(0)}</table>", result, flags=re.DOTALL)
    return result


def _inline(text: str) -> str:
    """Convert inline MD (bold, italic, code) to HTML."""
    import html as _html
    text = _html.escape(text)
    text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
    text = re.sub(r"\*(.+?)\*",     r"<em>\1</em>",         text)
    text = re.sub(r"`(.+?)`",       r"<code>\1</code>",     text)
    return text


# ─────────────────────────────────────────────────────────────────────────────
# Preview — assemble sections from DB into markdown, for on-screen rendering
# ─────────────────────────────────────────────────────────────────────────────

def assemble_preview(job_id: str) -> dict:
    """
    Return the full document content assembled from the DB — no file I/O.

    Designed for the GET /api/generate/{job_id}/preview endpoint so the
    frontend can render the document on-screen without downloading a file.

    Returns:
        {
          job_id, status, document_type, project_name,
          total_sections, completed_sections,
          sections: [{order, title, content, status, word_count}],
          markdown: "<full assembled markdown string>",
          export_urls: {docx, pdf, markdown},   # relative paths for the API
          blob_url: str | None,                 # GCS URL (gs://...) if in cloud mode + already exported
        }
    """
    from generation.db import GenerationJob, Section, get_session

    with get_session() as session:
        job = session.get(GenerationJob, job_id)
        if not job:
            raise ValueError(f"Job '{job_id}' not found")

        project_name = _extract_project_name(job.user_inputs_json)
        doc_type     = job.document_type or "Document"

        sections_out = []
        for sec in sorted(job.sections, key=lambda s: s.order_index):
            accepted = next((v for v in sec.versions if v.is_accepted), None)
            latest   = max(sec.versions, key=lambda v: v.version_number) if sec.versions else None
            chosen   = accepted or latest
            sections_out.append({
                "order":      sec.order_index,
                "section_id": sec.section_id,
                "title":      sec.section_title,
                "content":    chosen.content if chosen else "",
                "status":     sec.status,
                "word_count": chosen.word_count if chosen else 0,
            })

        status             = job.status
        total_sections     = job.total_sections
        completed_sections = job.completed_sections

    # Assemble full markdown
    lines = [
        f"# {doc_type}",
        f"**Project:** {project_name}",
        f"**Generated:** {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}",
        "",
        "---",
        "",
    ]
    for sec in sections_out:
        if sec["content"]:
            lines.append(f"## {sec['title']}")
            lines.append("")
            lines.append(sec["content"])
            lines.append("")
            lines.append("---")
            lines.append("")
    markdown = "\n".join(lines)

    # Blob URL: if a cloud-exported file already exists in local_storage (for local),
    # or the GCS URL was stored (cloud mode). We check for the most recent
    # exported file in the output dir and derive the URL.
    blob_url = _get_existing_blob_url(job_id)

    return {
        "job_id":             job_id,
        "status":             status,
        "document_type":      doc_type,
        "project_name":       project_name,
        "total_sections":     total_sections,
        "completed_sections": completed_sections,
        "sections":           sections_out,
        "markdown":           markdown,
        "export_urls": {
            "docx":     f"/api/generate/{job_id}/export?format=docx",
            "pdf":      f"/api/generate/{job_id}/export?format=pdf",
            "markdown": f"/api/generate/{job_id}/export?format=md",
        },
        "blob_url": blob_url,
    }


def export_job_to_temp(job_id: str, out_dir: Path) -> Path:
    """
    Export a job's sections as a DOCX file into an existing directory.
    Used by preview_service.py to feed LibreOffice headless conversion.

    Returns the Path to the written .docx file.
    Raises ValueError if no sections exist yet.
    """
    from generation.db import GenerationJob, get_session

    with get_session() as session:
        job = session.get(GenerationJob, job_id)
        if not job:
            raise ValueError(f"Job {job_id} not found")

        doc_type     = job.document_type
        project_name = _extract_project_name(job.user_inputs_json)

        sections_content: list[dict] = []
        for sec in sorted(job.sections, key=lambda s: s.order_index):
            if not sec.versions:
                continue
            accepted = next((v for v in sec.versions if v.is_accepted), None)
            latest   = max(sec.versions, key=lambda v: v.version_number)
            chosen   = accepted or latest
            sections_content.append({
                "key":     getattr(sec, "section_key", None) or "",
                "title":   sec.section_title,
                "content": chosen.content,
                "version": chosen.version_number,
            })

    if not sections_content:
        raise ValueError("No completed sections to export")

    safe_name = re.sub(r"[^\w\s-]", "", project_name)[:40].strip().replace(" ", "_") or "document"
    out_path  = out_dir / f"{safe_name}_preview.docx"

    _write_styled_docx(out_path, doc_type, project_name, sections_content)

    return out_path


def upload_output_to_blob(job_id: str, file_path: Path, mime_type: str) -> Optional[str]:
    """
    Upload an exported output file to GCS.
    Returns the GCS URL (gs://...), or None if running in local mode or upload fails.
    Only called when LOCAL_MODE=false (GCS storage mode).
    """
    if LOCAL_DB:
        return None
    try:
        from storage.gcs_storage import get_storage_service, GCSStorageService
        store = get_storage_service()
        if not isinstance(store, GCSStorageService):
            return None
        blob_path = f"outputs/{_safe_job_id(job_id)}/{file_path.name}"
        blob_url = store._upload_blob(blob_path, file_path.read_bytes(), mime_type)
        logger.info("[preview] Uploaded output to GCS: %s", blob_url)
        return blob_url
    except Exception as e:
        logger.warning("[preview] GCS upload failed (non-fatal): %s", e)
        return None


def _get_existing_blob_url(job_id: str) -> Optional[str]:
    """
    In cloud mode: list GCS objects in outputs/{job_id}/ and return the DOCX URL if found.
    In local mode: return None (frontend uses the export endpoint instead).
    """
    if LOCAL_DB:
        return None
    try:
        from storage.gcs_storage import get_storage_service, GCSStorageService
        store = get_storage_service()
        if not isinstance(store, GCSStorageService):
            return None
        # List blobs under outputs/{job_id}/
        blobs = list(store._client.list_blobs(store._bucket_name, prefix=f"outputs/{_safe_job_id(job_id)}/"))
        # Prefer DOCX, then PDF, then markdown
        for ext in (".docx", ".pdf", ".md"):
            for b in blobs:
                if b.name.endswith(ext):
                    return f"gs://{store._bucket_name}/{b.name}"
    except Exception as e:
        logger.debug("[preview] Could not check GCS URL: %s", e)
    return None


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _safe_job_id(job_id: str) -> str:
    """Validate a job_id before using it in a filesystem/blob path. Real job_ids
    are UUIDs; reject anything containing path separators or traversal sequences
    so a crafted id can never escape the outputs directory (CWE-22)."""
    import re as _re
    jid = (job_id or "").strip()
    if not _re.fullmatch(r"[A-Za-z0-9_-]{1,64}", jid):
        raise ValueError(f"Invalid job_id: {job_id!r}")
    return jid


def _output_dir(job_id: str) -> Path:
    # Absolute path — works regardless of the working directory at runtime.
    # job_id is validated (no separators / '..') to prevent path traversal.
    return Path(__file__).parent.parent / "local_storage" / "outputs" / _safe_job_id(job_id)


def _is_brd(doc_type: Optional[str]) -> bool:
    if not doc_type:
        return False
    return "brd" in doc_type.lower() or "business requirements" in doc_type.lower()


def _is_nit(doc_type: Optional[str]) -> bool:
    if not doc_type:
        return False
    dt = doc_type.lower()
    return "nit" in dt or "notice inviting tender" in dt


def _extract_project_name(user_inputs_json: Optional[str]) -> str:
    if not user_inputs_json:
        return "Document"
    try:
        import json
        data = json.loads(user_inputs_json)
        return data.get("project_name") or "Document"
    except Exception:
        return "Document"
